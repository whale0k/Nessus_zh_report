#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word generator module - Generate comprehensive vulnerability assessment Word reports
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.table import _Cell
from datetime import datetime
from typing import List, Dict, Any
import logging
from collections import defaultdict
import matplotlib.pyplot as plt
import io
import os
import tempfile

from config.constants import AppConstants
from models import VulnerabilityInfo
from utils import RiskCalculator, StatisticsCalculator

logger = logging.getLogger(__name__)


class WordReportGenerator:
    """Generate comprehensive vulnerability assessment Word reports"""

    def __init__(self, template_path: str = None):
        self.template_path = template_path or "./config/漏洞扫描汇报及修复建议_漏洞排序模板.docx"
        self.risk_calculator = RiskCalculator()
        self.stats_calculator = StatisticsCalculator()

    @staticmethod
    def translate_risk_level(risk_level: str) -> str:
        """Translate English risk levels to Chinese"""
        return AppConstants.RISK_LEVEL_MAPPING.get(risk_level, risk_level)

    def create_vulnerability_report(
        self,
        vulnerabilities: List[VulnerabilityInfo],
        config: Dict[str, Any] = None,
        output_file: str = None,
        filter_level: str = "all",
        risk_counts: Dict[str, int] = None,
        host_risk_counts: Dict[str, int] = None,
        total_hosts: int = None
    ) -> str:
        """Create comprehensive vulnerability assessment report"""
        logger.info(f"Creating Word report, filter_level: {filter_level}")

        if not output_file:
            current_date = datetime.now().strftime("%Y_%m")
            output_file = f"漏洞扫描汇报及修复建议({current_date}).docx"

        if filter_level == "high_above":
            high_risk_levels = ["Critical", "High"]
            filtered_vulnerabilities = [v for v in vulnerabilities if v.risk in high_risk_levels]
            logger.info(
                f"Word report: Filtered to {len(filtered_vulnerabilities)} "
                f"high-risk vulnerabilities (from {len(vulnerabilities)} total)"
            )
        else:
            filtered_vulnerabilities = vulnerabilities
            logger.info(f"Word report: Including all {len(vulnerabilities)} vulnerabilities")

        try:
            doc = Document(self.template_path)
            logger.info(f"Loaded template: {self.template_path}")

            stats = self._calculate_statistics(
                vulnerabilities, risk_counts, host_risk_counts, total_hosts
            )

            default_config = {
                "客户名称": "目标企业",
                "实施公司": "安全服务公司"
            }
            if config:
                default_config.update(config)

            self._replace_placeholders(doc, stats, default_config)

            self._add_vulnerability_details(doc, filtered_vulnerabilities, stats)

            if filter_level == "high_above":
                logger.info("Calling _remove_unused_risk_sections")
                self._remove_unused_risk_sections(doc)
                logger.info("_remove_unused_risk_sections completed")

            doc.save(output_file)
            logger.info(f"Word report saved: {output_file}")

            if filter_level == "high_above":
                logger.info("Attempting to update table of contents using Word COM...")
                success = self._update_table_of_contents_with_word(output_file)
                if success:
                    logger.info("Table of contents updated successfully")
                else:
                    logger.warning(
                        "Could not auto-update TOC. Please update manually in Word (Ctrl+A then F9)"
                    )

            return output_file

        except Exception as e:
            logger.error(f"Error creating Word report: {e}")
            raise

    def _calculate_statistics(
        self,
        vulnerabilities: List[VulnerabilityInfo],
        risk_counts: Dict[str, int] = None,
        host_risk_counts: Dict[str, int] = None,
        total_hosts: int = None
    ) -> Dict[str, Any]:
        """Calculate comprehensive vulnerability and host statistics"""
        if risk_counts is not None and host_risk_counts is not None and total_hosts is not None:
            logger.info("Using pre-calculated statistics")

            total_vulns = len(vulnerabilities)

            def safe_percentage(count, total):
                return f"{count/total*100:.1f}%" if total > 0 else "0.0%"

            return {
                "主机数总计": total_hosts,
                "超危主机-数量": host_risk_counts.get("超危", 0),
                "超危主机-占比": safe_percentage(host_risk_counts.get("超危", 0), total_hosts),
                "高危主机-数量": host_risk_counts.get("高危", 0),
                "高危主机-占比": safe_percentage(host_risk_counts.get("高危", 0), total_hosts),
                "中危主机-数量": host_risk_counts.get("中危", 0),
                "中危主机-占比": safe_percentage(host_risk_counts.get("中危", 0), total_hosts),
                "低危主机-数量": host_risk_counts.get("低危", 0),
                "低危主机-占比": safe_percentage(host_risk_counts.get("低危", 0), total_hosts),
                "安全主机-数量": host_risk_counts.get("安全", 0),
                "安全主机-占比": safe_percentage(host_risk_counts.get("安全", 0), total_hosts),
                "中危以上主机-占比": safe_percentage(
                    host_risk_counts.get("超危", 0) + host_risk_counts.get("高危", 0) +
                    host_risk_counts.get("中危", 0),
                    total_hosts
                ),
                "漏洞数量总计": total_vulns,
                "超危漏洞-数量": risk_counts.get("Critical", 0),
                "超危漏洞-占比": safe_percentage(risk_counts.get("Critical", 0), total_vulns),
                "高危漏洞-数量": risk_counts.get("High", 0),
                "高危漏洞-占比": safe_percentage(risk_counts.get("High", 0), total_vulns),
                "中危漏洞-数量": risk_counts.get("Medium", 0),
                "中危漏洞-占比": safe_percentage(risk_counts.get("Medium", 0), total_vulns),
                "低危漏洞-数量": risk_counts.get("Low", 0),
                "低危漏洞-占比": safe_percentage(risk_counts.get("Low", 0), total_vulns),
                "信息漏洞-数量": risk_counts.get("Info", 0),
                "信息漏洞-占比": safe_percentage(risk_counts.get("Info", 0), total_vulns),
                "vulnerabilities_by_risk": self.stats_calculator.group_vulnerabilities_by_risk(vulnerabilities)
            }

        logger.info("Using original statistics calculation logic")

        host_vulnerabilities = defaultdict(list)
        for vuln in vulnerabilities:
            host_vulnerabilities[vuln.host].append(vuln)

        host_risks = {}
        for host, vulns in host_vulnerabilities.items():
            cvss_scores = [v.cvss_v3_score or v.cvss_v2_score for v in vulns
                          if v.cvss_v3_score or v.cvss_v2_score]
            risk_score, risk_level = self.risk_calculator.calculate_host_risk(cvss_scores)
            host_risks[host] = {
                'risk_score': risk_score,
                'risk_level': risk_level,
                'vulnerabilities': vulns
            }

        vuln_risk_counts = defaultdict(int)
        for vuln in vulnerabilities:
            risk_level = self.stats_calculator.normalize_risk_level(vuln.risk)
            vuln_risk_counts[risk_level] += 1

        host_risk_counts_computed = defaultdict(int)
        for host_info in host_risks.values():
            host_risk_counts_computed[host_info['risk_level']] += 1

        total_hosts_computed = len(host_risks)
        total_vulns = len(vulnerabilities)

        def safe_percentage(count, total):
            return f"{count/total*100:.1f}%" if total > 0 else "0.0%"

        return self.stats_calculator.calculate_comprehensive_stats(vulnerabilities, host_risks)

    def _create_host_risk_chart(self, stats: Dict[str, Any], output_path: str) -> str:
        """Create host risk distribution pie chart"""
        try:
            plt.rcParams['font.sans-serif'] = AppConstants.CHINESE_FONTS
            plt.rcParams['axes.unicode_minus'] = False

            labels = []
            sizes = []
            chart_colors = []

            risk_levels = [
                ('超危主机-数量', '超危'),
                ('高危主机-数量', '高危'),
                ('中危主机-数量', '中危'),
                ('低危主机-数量', '低危'),
                ('安全主机-数量', '安全')
            ]

            for key, label in risk_levels:
                count = stats.get(key, 0)
                if count > 0:
                    labels.append(f'{label}({count}台)')
                    sizes.append(count)
                    chart_colors.append(AppConstants.CHART_COLORS[label])

            if not sizes:
                labels = ['无数据']
                sizes = [1]
                chart_colors = ['#CCCCCC']

            fig, ax = plt.subplots(figsize=(8, 6))

            wedges, texts, autotexts = ax.pie(
                sizes, labels=labels, autopct='%1.1f%%', colors=chart_colors
            )

            ax.set_title('主机风险分布', fontsize=16, fontweight='bold', pad=20)
            ax.axis('equal')

            plt.tight_layout()
            plt.savefig(output_path, dpi=300, bbox_inches='tight',
                       facecolor='white', edgecolor='none')
            plt.close()

            logger.info(f"Host risk chart saved: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error creating host risk chart: {e}")
            fig, ax = plt.subplots(figsize=(8, 6))
            ax.text(0.5, 0.5, '图表生成失败', ha='center', va='center', fontsize=16)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
            ax.axis('off')
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()
            return output_path

    def _create_vulnerability_risk_chart(self, stats: Dict[str, Any], output_path: str) -> str:
        """Create vulnerability risk distribution pie chart"""
        try:
            plt.rcParams['font.sans-serif'] = AppConstants.CHINESE_FONTS
            plt.rcParams['axes.unicode_minus'] = False

            labels = []
            sizes = []
            chart_colors = []

            risk_levels = [
                ('超危漏洞-数量', '超危'),
                ('高危漏洞-数量', '高危'),
                ('中危漏洞-数量', '中危'),
                ('低危漏洞-数量', '低危'),
                ('信息漏洞-数量', '信息')
            ]

            for key, label in risk_levels:
                count = stats.get(key, 0)
                if count > 0:
                    labels.append(f'{label}({count}个)')
                    sizes.append(count)
                    chart_colors.append(AppConstants.CHART_COLORS[label])

            if not sizes:
                labels = ['无数据']
                sizes = [1]
                chart_colors = ['#CCCCCC']

            fig, ax = plt.subplots(figsize=(8, 6))

            wedges, texts, autotexts = ax.pie(
                sizes, labels=labels, autopct='%1.1f%%', colors=chart_colors
            )

            ax.set_title('漏洞风险分布', fontsize=16, fontweight='bold', pad=20)
            ax.axis('equal')

            plt.tight_layout()
            plt.savefig(output_path, dpi=300, bbox_inches='tight',
                       facecolor='white', edgecolor='none')
            plt.close()

            logger.info(f"Vulnerability risk chart saved: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error creating vulnerability risk chart: {e}")
            fig, ax = plt.subplots(figsize=(8, 6))
            ax.text(0.5, 0.5, '图表生成失败', ha='center', va='center', fontsize=16)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
            ax.axis('off')
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()
            return output_path

    def _replace_placeholders(
        self,
        doc: Document,
        stats: Dict[str, Any],
        config: Dict[str, Any]
    ) -> None:
        """Replace placeholders in the document"""
        current_time = datetime.now()
        current_year = str(current_time.year)
        current_month = str(current_time.month)

        replacements = {
            "{xxxx客户名称}": config.get("客户名称", ""),
            "{xxxx实施公司}": config.get("实施公司", ""),
            "{时间-年}": current_year,
            "{时间-月}": current_month,
            "{主机数总计}": str(stats["主机数总计"]),
            "{超危主机-数量}": str(stats["超危主机-数量"]),
            "{超危主机-占比}": stats["超危主机-占比"],
            "{高危主机-数量}": str(stats["高危主机-数量"]),
            "{高危主机-占比}": stats["高危主机-占比"],
            "{中危主机-数量}": str(stats["中危主机-数量"]),
            "{中危主机-占比}": stats["中危主机-占比"],
            "{低危主机-数量}": str(stats["低危主机-数量"]),
            "{低危主机-占比}": stats["低危主机-占比"],
            "{安全主机-数量}": str(stats["安全主机-数量"]),
            "{安全主机-占比}": stats["安全主机-占比"],
            "{中危以上主机-占比}": stats["中危以上主机-占比"],
            "{超危漏洞-数量}": str(stats["超危漏洞-数量"]),
            "{超危漏洞-占比}": stats["超危漏洞-占比"],
            "{高危漏洞-数量}": str(stats["高危漏洞-数量"]),
            "{高危漏洞-占比}": stats["高危漏洞-占比"],
            "{中危漏洞-数量}": str(stats["中危漏洞-数量"]),
            "{中危漏洞-占比}": stats["中危漏洞-占比"],
            "{低危漏洞-数量}": str(stats["低危漏洞-数量"]),
            "{低危漏洞-占比}": stats["低危漏洞-占比"],
            "{信息漏洞-数量}": str(stats["信息漏洞-数量"]),
            "{信息漏洞-占比}": stats["信息漏洞-占比"],
            "{漏洞数量总计}": str(stats["漏洞数量总计"])
        }

        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
                    logger.debug(f"Replaced {placeholder} with {replacement}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, replacement)

        self._replace_chart_placeholders(doc, stats)

    def _replace_placeholder_with_image(
        self,
        doc: Document,
        placeholder: str,
        image_path: str,
        width: Inches = None
    ) -> None:
        """Replace placeholder with image"""
        try:
            if width is None:
                width = Inches(6)

            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, "")
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(image_path, width=width)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logger.info(f"Replaced {placeholder} with image")
                    return

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, "")
                                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                                run.add_picture(image_path, width=width)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                logger.info(f"Replaced {placeholder} with image (in table)")
                                return

            logger.warning(f"Placeholder {placeholder} not found")

        except Exception as e:
            logger.error(f"Error replacing placeholder {placeholder} with image: {e}")

    def _replace_chart_placeholders(
        self,
        doc: Document,
        stats: Dict[str, Any]
    ) -> None:
        """Replace chart placeholders with generated images"""
        try:
            temp_dir = tempfile.mkdtemp()

            chart_placeholders = {
                "{主机风险分布图}": "host_risk_chart.png",
                "{漏洞风险分布图}": "vulnerability_risk_chart.png"
            }

            for placeholder, filename in chart_placeholders.items():
                chart_path = os.path.join(temp_dir, filename)

                if "主机风险" in placeholder:
                    self._create_host_risk_chart(stats, chart_path)
                elif "漏洞风险" in placeholder:
                    self._create_vulnerability_risk_chart(stats, chart_path)

                self._replace_placeholder_with_image(doc, placeholder, chart_path)

        except Exception as e:
            logger.error(f"Error processing chart placeholders: {e}")

    def _add_vulnerability_details(
        self,
        doc: Document,
        vulnerabilities: List[VulnerabilityInfo],
        stats: Dict[str, Any]
    ) -> None:
        """Add detailed vulnerability information tables"""
        vulnerabilities_by_risk = stats["vulnerabilities_by_risk"]

        risk_levels = [
            ("Critical", "严重漏洞详情"),
            ("High", "高危漏洞详情"),
            ("Medium", "中危漏洞详情"),
            ("Low", "低危漏洞详情")
        ]

        for risk_level, placeholder in risk_levels:
            placeholder_text = f"{{{placeholder}}}"
            if risk_level in vulnerabilities_by_risk:
                vulns = vulnerabilities_by_risk[risk_level]
                detail_tables = self._create_vulnerability_detail_tables(vulns)
                self._find_and_replace_placeholder(doc, placeholder_text, detail_tables)
            else:
                self._find_and_replace_placeholder(doc, placeholder_text, [])

    def _create_vulnerability_detail_tables(
        self,
        vulnerabilities: List[VulnerabilityInfo]
    ) -> List[Dict[str, str]]:
        """Create detailed vulnerability tables"""
        tables_data = []

        vuln_groups = defaultdict(list)
        for vuln in vulnerabilities:
            vuln_groups[vuln.name].append(vuln)

        for vuln_name, vuln_list in vuln_groups.items():
            repr_vuln = vuln_list[0]

            affected_hosts = [v.host for v in vuln_list]
            affected_hosts = list(set(affected_hosts))

            affected_ports = [v.port for v in vuln_list if v.port]
            affected_ports = list(set(affected_ports))

            table_data = {
                "漏洞名称": repr_vuln.name,
                "主机": "，".join(affected_hosts),
                "详细描述": repr_vuln.description,
                "漏洞端口": "，".join(affected_ports) if affected_ports else "N/A",
                "漏洞等级": self.translate_risk_level(repr_vuln.risk),
                "CVE编号": repr_vuln.cve or "N/A",
                "修补建议": repr_vuln.solution
            }

            tables_data.append(table_data)

        return tables_data

    def _find_and_replace_placeholder(
        self,
        doc: Document,
        placeholder: str,
        tables_data: List[Dict[str, str]]
    ) -> None:
        """Find placeholder and replace with vulnerability tables"""
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                self._replace_paragraph_with_table(doc, paragraph, tables_data, placeholder)
                return

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_paragraph_with_table(doc, paragraph, tables_data, placeholder)
                            return

    def _replace_paragraph_with_table(
        self,
        doc: Document,
        paragraph,
        tables_data: List[Dict[str, str]],
        placeholder: str
    ) -> None:
        """Replace specific paragraph with vulnerability table"""
        parent = paragraph._element.getparent()
        paragraph_index = list(parent).index(paragraph._element)

        paragraph.text = paragraph.text.replace(placeholder, "")

        if tables_data:
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'

            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(4.5)

            for i, vuln_data in enumerate(tables_data):
                name_row = table.add_row()
                name_cell = name_row.cells[0]
                name_cell.merge(name_row.cells[1])

                name_paragraph = name_cell.paragraphs[0]
                name_paragraph.clear()
                name_run = name_paragraph.add_run(vuln_data["漏洞名称"])
                name_run.font.bold = True

                self._set_cell_background_color(name_cell, AppConstants.CELL_BACKGROUND_COLOR)
                name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for key, value in vuln_data.items():
                    if key != "漏洞名称":
                        detail_row = table.add_row()
                        detail_row.cells[0].text = key
                        detail_row.cells[1].text = value

                        self._set_cell_background_color(
                            detail_row.cells[0], AppConstants.CELL_BACKGROUND_COLOR
                        )
                        detail_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            table_element = table._element
            parent.insert(paragraph_index + 1, table_element)

            p = paragraph._element
            p.getparent().remove(p)

    @staticmethod
    def _set_cell_background_color(cell: _Cell, color: str) -> None:
        """Set cell background color"""
        try:
            cell_xml_element = cell._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), color)
            table_cell_properties.append(shade_obj)
        except Exception as e:
            logger.warning(f"Could not set cell background color: {e}")

    def _remove_unused_risk_sections(self, doc: Document) -> None:
        """Remove unused risk level section titles when filter is applied"""
        try:
            logger.info("Calling _remove_unused_risk_sections")
            target_texts = ["中危漏洞", "低危漏洞"]
            total_removed = 0

            for target_text in target_texts:
                logger.info(f"Looking for and deleting '{target_text}' sections...")
                paragraphs_to_remove = []

                for i, paragraph in enumerate(doc.paragraphs):
                    logger.debug(f"Checking paragraph {i}: style='{paragraph.style.name}', text='{paragraph.text}'")
                    if paragraph.style.name.startswith('Heading 2') and target_text in paragraph.text:
                        paragraphs_to_remove.append(paragraph)
                        logger.info(f"Found target heading: '{paragraph.text}' (position: {i})")

                if paragraphs_to_remove:
                    for paragraph in paragraphs_to_remove:
                        p = paragraph._element
                        p.getparent().remove(p)
                    logger.info(
                        f"Successfully deleted {len(paragraphs_to_remove)} paragraphs "
                        f"containing '{target_text}'"
                    )
                    total_removed += len(paragraphs_to_remove)
                else:
                    logger.info(f"No Heading 2 paragraphs found containing '{target_text}'")

            self._remove_toc_entries(doc, target_texts)

            logger.info(f"Total deleted {total_removed} paragraphs")
            doc._body._element = doc._body._element

        except Exception as e:
            logger.error(f"Error removing unused risk sections: {e}", exc_info=True)

    def _remove_toc_entries(self, doc: Document, target_texts: List[str]) -> None:
        """Delete TOC entries for removed sections"""
        try:
            logger.info("Deleting TOC entries for removed sections...")
            toc_entries_removed = 0

            for i, paragraph in enumerate(doc.paragraphs):
                if (paragraph.style.name.startswith('toc') or
                        'TOC' in paragraph.style.name or
                        any(target_text in paragraph.text for target_text in target_texts)):

                    for target_text in target_texts:
                        if target_text in paragraph.text:
                            logger.info(f"Deleting TOC entry: '{paragraph.text}'")
                            p = paragraph._element
                            p.getparent().remove(p)
                            toc_entries_removed += 1
                            break

            logger.info(f"Deleted {toc_entries_removed} TOC entries")
        except Exception as e:
            logger.error(f"Error deleting TOC entries: {e}", exc_info=True)

    @staticmethod
    def _update_table_of_contents_with_word(file_path: str) -> bool:
        """Update table of contents in Word document using win32com.client"""
        word_app = None
        try:
            import win32com.client

            abs_file_path = os.path.abspath(file_path)
            logger.info(f"Updating table of contents using Word COM: {abs_file_path}")

            if not os.path.exists(abs_file_path):
                logger.error(f"File not found: {abs_file_path}")
                return False

            try:
                word_app = win32com.client.Dispatch("Word.Application")
            except Exception:
                word_app = win32com.client.DispatchEx("kwps.Application")
            finally:
                if not word_app:
                    logger.warning("TOC update failed. Please install Word or WPS 2019 on Windows")

            word_app.Visible = False

            doc = word_app.Documents.Open(abs_file_path)

            doc.Fields.Update()

            doc.Save()
            doc.Close()

            logger.info("Table of contents updated successfully using Word COM")
            return True

        except ImportError:
            logger.warning(
                "win32com.client not available. Install with: pip install pywin32"
            )
            return False
        except Exception as e:
            logger.error(f"Error updating table of contents using Word COM: {e}")
            return False
        finally:
            if word_app:
                try:
                    word_app.Quit()
                except Exception as e:
                    logger.warning(f"Error closing Word application: {e}")
